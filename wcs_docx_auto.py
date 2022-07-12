from docx import Document
from docx.shared import Inches
import docx
import re
import datetime

def now_date():
    now = datetime.datetime.now()
    now_dt = now.strftime('%Y-%m-%d')
    result = re.sub('[^0-9]', '', now_dt)
    print('today',result)
    return result

try:
    document = Document('C:/Users/ljs91/Desktop/Olleh_TV_홈포털_플랫폼_운영보고서_' + now_date() + '.docx')

    file_list = ['C:/Users/ljs91/Desktop/WCS_정기점검_샘플데이터/capture/이미지 045.png',
                 'C:/Users/ljs91/Desktop/WCS_정기점검_샘플데이터/capture/이미지 046.png']
    file_list1 = ['C:/Users/ljs91/Desktop/WCS_정기점검_샘플데이터/capture/ocr대상/이미지 043.png',
                  'C:/Users/ljs91/Desktop/WCS_정기점검_샘플데이터/capture/ocr대상/이미지 044.png']

    tables = document.tables
    num = 6
    for i in file_list:
        p = tables[num].rows[0].cells[0].add_paragraph()
        r = p.add_run()
        r.add_picture(i, width=Inches(6.7))
        num += 1
        if num == 8:
            for k in file_list1:
                p = tables[num].rows[0].cells[0].add_paragraph()
                r = p.add_run()
                r.add_picture(k, width=Inches(6.7))
                num += 1
                if num == 10: break
    document.save('C:/Users/ljs91/Desktop/Olleh_TV_홈포털_플랫폼_운영보고서_' + now_date() + '.docx')

except NameError as e:
    print('Error', e)

except docx.opc.exceptions.PackageNotFoundError as e:
    print('Error', e)
